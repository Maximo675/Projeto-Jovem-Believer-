"""
Script para converter arquivos .docx para .md
Usa os documentos INFANT.ID para recriar as aulas
"""

from docx import Document
from pathlib import Path
import os

docs_folder = Path("assets/documents")

def convert_docx_to_md(docx_path, md_path):
    """Converte um arquivo .docx para .md preservando estrutura"""
    
    doc = Document(docx_path)
    md_content = ""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            md_content += "\n"
            continue
        
        # Detectar nível de título pela formatação
        if para.style.name.startswith('Heading'):
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
            md_content += f"{'#' * level} {text}\n\n"
        else:
            md_content += f"{text}\n"
    
    # Processar tabelas
    for table in doc.tables:
        md_content += "\n| "
        for cell in table.rows[0].cells:
            md_content += f"{cell.text} | "
        md_content += "\n|"
        for _ in table.rows[0].cells:
            md_content += " --- |"
        md_content += "\n"
        
        for row in table.rows[1:]:
            md_content += "| "
            for cell in row.cells:
                md_content += f"{cell.text} | "
            md_content += "\n"
        
        md_content += "\n"
    
    # Salvar como markdown
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    return True

# Converter todos os .docx
print("=" * 60)
print("Convertendo arquivos .docx para .md")
print("=" * 60)

converted = 0
for docx_file in docs_folder.glob("*.docx"):
    md_file = docx_file.with_suffix(".md")
    try:
        convert_docx_to_md(docx_file, md_file)
        print(f"✓ {docx_file.name}")
        print(f"  → {md_file.name}")
        converted += 1
    except Exception as e:
        print(f"✗ Erro ao converter {docx_file.name}: {e}")

print("=" * 60)
print(f"Conversao completa! [{converted} arquivo(s) convertido(s)]")
print("=" * 60)

# Listar arquivos .md criados
print("\nArquivos .md criados:")
for md_file in docs_folder.glob("*.md"):
    size = os.path.getsize(md_file) / 1024
    print(f"  • {md_file.name} ({size:.1f} KB)")
