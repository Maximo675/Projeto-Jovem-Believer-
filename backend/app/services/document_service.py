"""
Serviço de Gerenciamento de Documentos
Processa e disponibiliza documentos Word como conteúdo educacional
"""

import os
from pathlib import Path
from datetime import datetime
from app import db
from app.models.course import Course
from app.models.lesson import Lesson

class DocumentService:
    """Serviço para processar documentos Word como base de conhecimento"""
    
    KNOWLEDGE_BASE_DIR = Path(__file__).parent.parent.parent / 'assets' / 'documents'
    
    @staticmethod
    def listar_documentos():
        """Lista todos os documentos disponíveis na knowledge base"""
        try:
            documentos = []
            if DocumentService.KNOWLEDGE_BASE_DIR.exists():
                for arquivo in DocumentService.KNOWLEDGE_BASE_DIR.glob('*.docx'):
                    documentos.append({
                        'nome': arquivo.stem,
                        'arquivo': arquivo.name,
                        'caminho': str(arquivo),
                        'tamanho': arquivo.stat().st_size,
                        'modificado': datetime.fromtimestamp(arquivo.stat().st_mtime).isoformat()
                    })
            return documentos
        except Exception as e:
            raise Exception(f"Erro ao listar documentos: {str(e)}")
    
    @staticmethod
    def obter_documento(nome_arquivo):
        """Obtém um documento específico"""
        try:
            caminho = DocumentService.KNOWLEDGE_BASE_DIR / nome_arquivo
            
            if not caminho.exists():
                raise FileNotFoundError(f"Documento {nome_arquivo} não encontrado")
            
            if not str(caminho).endswith('.docx'):
                raise ValueError("Formato de arquivo não suportado")
            
            return {
                'nome': caminho.stem,
                'arquivo': caminho.name,
                'caminho': str(caminho),
                'tamanho': caminho.stat().st_size,
                'conteudo': DocumentService.extrair_conteudo(caminho)
            }
        except Exception as e:
            raise Exception(f"Erro ao obter documento: {str(e)}")
    
    @staticmethod
    def extrair_conteudo(caminho_arquivo):
        """
        Extrai conteúdo de um arquivo Word
        Usa python-docx para extrair texto real
        """
        try:
            from docx import Document
            
            doc = Document(caminho_arquivo)
            
            # Extrair parágrafos
            paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extrair tabelas
            tabelas_texto = []
            for tabela in doc.tables:
                for linha in tabela.rows:
                    linha_texto = [celula.text for celula in linha.cells]
                    tabelas_texto.append(' | '.join(linha_texto))
            
            # Combinar conteúdo
            conteudo = '\n'.join(paragrafos)
            if tabelas_texto:
                conteudo += '\n\n[TABELAS]\n' + '\n'.join(tabelas_texto)
            
            return conteudo
            
        except ImportError:
            return "Erro: python-docx não está instalado. Execute: pip install python-docx"
        except Exception as e:
            raise Exception(f"Erro ao extrair conteúdo: {str(e)}")
    
    @staticmethod
    def criar_curso_de_documento(nome_documento, titulo_curso, descricao, nivel='basico'):
        """
        Cria um curso baseado em um documento Word
        Útil para transformar documentos em cursos estruturados
        """
        try:
            documento = DocumentService.obter_documento(f"{nome_documento}.docx")
            
            # Criar curso
            curso = Course(
                titulo=titulo_curso,
                descricao=descricao,
                nivel=nivel,
                autor='INFANT.ID System',
                imagem_url=None
            )
            
            db.session.add(curso)
            db.session.commit()
            
            # Criar aula com conteúdo do documento
            aula = Lesson(
                curso_id=curso.id,
                titulo=nome_documento,
                descricao=f"Conteúdo extraído de: {documento['arquivo']}",
                conteudo=DocumentService.formatar_conteudo_html(documento['conteudo']),
                ordem=1,
                duracao=None
            )
            
            db.session.add(aula)
            db.session.commit()
            
            return {
                'curso_id': curso.id,
                'curso_titulo': curso.titulo,
                'aula_id': aula.id,
                'mensagem': 'Curso criado com sucesso!'
            }
        except Exception as e:
            db.session.rollback()
            raise Exception(f"Erro ao criar curso: {str(e)}")
    
    @staticmethod
    def formatar_conteudo_html(conteudo):
        """Formata conteúdo para HTML"""
        return f"""
        <div class="conteudo-documento">
            <h2>Conteúdo do Documento</h2>
            <p>{conteudo}</p>
            <div class="nota-importante">
                <strong>Nota:</strong> Este conteúdo foi extraído de um documento fornecido.
                Para uma melhor experiência, a versão completa do documento está disponível para download.
            </div>
        </div>
        """


class KnowledgeBaseManager:
    """Gerenciador da base de conhecimento da plataforma"""
    
    @staticmethod
    def sincronizar_documentos():
        """Sincroniza documentos da pasta com o banco de dados"""
        try:
            documentos = DocumentService.listar_documentos()
            
            return {
                'total': len(documentos),
                'documentos': documentos,
                'status': 'sincronizado'
            }
        except Exception as e:
            return {
                'total': 0,
                'documentos': [],
                'status': 'erro',
                'mensagem': str(e)
            }
    
    @staticmethod
    def gerar_indice_conhecimento():
        """Gera índice de todos os documentos e cursos"""
        try:
            documentos = DocumentService.listar_documentos()
            cursos = Course.query.all()
            
            indice = {
                'data_geracao': datetime.utcnow().isoformat(),
                'total_documentos': len(documentos),
                'total_cursos': len(cursos),
                'documentos': documentos,
                'cursos': [c.to_dict() for c in cursos]
            }
            
            return indice
        except Exception as e:
            raise Exception(f"Erro ao gerar índice: {str(e)}")
