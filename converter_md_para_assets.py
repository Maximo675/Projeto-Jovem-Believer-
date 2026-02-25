#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Converter Markdown para DOCX
E colocar em assets/documents/ para integração com o site
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def converter_md_para_docx(md_path, docx_path):
    """Converte Markdown para Word (.docx)"""
    
    # Ler arquivo markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    
    # Criar documento Word
    doc = Document()
    
    # Processar linha por linha
    linhas = conteudo.split('\n')
    
    for linha in linhas:
        linha_limpa = linha.strip()
        
        if not linha_limpa:
            doc.add_paragraph()
            continue
        
        # Títulos principais (# Título)
        if linha_limpa.startswith('# '):
            titulo = linha_limpa[2:].strip()
            p = doc.add_paragraph(titulo, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # Títulos secundários (## Título)
        if linha_limpa.startswith('## '):
            titulo = linha_limpa[3:].strip()
            doc.add_paragraph(titulo, style='Heading 2')
            continue
        
        # Títulos terciários (### Título)
        if linha_limpa.startswith('### '):
            titulo = linha_limpa[4:].strip()
            doc.add_paragraph(titulo, style='Heading 3')
            continue
        
        # Listas com "-" ou "*"
        if linha_limpa.startswith('- ') or linha_limpa.startswith('* '):
            item = linha_limpa[2:].strip()
            doc.add_paragraph(item, style='List Bullet')
            continue
        
        # Listas numeradas
        if re.match(r'^\d+\. ', linha_limpa):
            item = re.sub(r'^\d+\. ', '', linha_limpa)
            doc.add_paragraph(item, style='List Number')
            continue
        
        # Paragrafos normais
        doc.add_paragraph(linha_limpa)
    
    # Salvar
    doc.save(docx_path)
    print(f"✅ Convertido para: {docx_path}")
    return docx_path

def main():
    # Caminhos
    md_file = Path(r"c:\Users\maximo.silva\Desktop\Desevolvimento\Alura Jovem Believer\notes\TREINAMENTO_REPLICADORES_ENFERMEIRAS.md")
    docs_dir = Path(r"c:\Users\maximo.silva\Desktop\Desevolvimento\Alura Jovem Believer\assets\documents")
    
    # Criar diretório se não existir
    docs_dir.mkdir(parents=True, exist_ok=True)
    
    # Caminhos de saída
    docx_file = docs_dir / "Treinamento Replicadores Enfermeiras.docx"
    md_file_assets = docs_dir / "Treinamento Replicadores Enfermeiras.md"
    
    print("🔄 Convertendo Markdown para Word...")
    print(f"📖 Origem: {md_file}")
    print(f"📁 Destino: {docs_dir}")
    print()
    
    # Converter para DOCX
    converter_md_para_docx(md_file, docx_file)
    
    # Copiar também o MD para assets (como fazem os outros documentos)
    with open(md_file, 'r', encoding='utf-8') as f:
        conteudo_md = f.read()
    
    with open(md_file_assets, 'w', encoding='utf-8') as f:
        f.write(conteudo_md)
    
    print(f"✅ Copiado Markdown para: {md_file_assets}")
    
    print()
    print("=" * 60)
    print("✅ INTEGRAÇÃO COMPLETA!")
    print("=" * 60)
    print()
    print("Arquivos criados em assets/documents/:")
    print(f"  ✅ Treinamento Replicadores Enfermeiras.docx ({docx_file.stat().st_size / 1024:.1f} KB)")
    print(f"  ✅ Treinamento Replicadores Enfermeiras.md ({md_file_assets.stat().st_size / 1024:.1f} KB)")
    print()
    print("🌐 Agora acessível via:")
    print("  → Dashboard: Seção 'Documentos'")
    print("  → API: GET /api/documents/Treinamento%20Replicadores%20Enfermeiras")
    print()

if __name__ == "__main__":
    main()
